import os
import pandas as pd
import pdfminer.high_level
import spacy
import json
import csv
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
from fuzzywuzzy import fuzz
from tabulate import tabulate
from openpyxl import Workbook
import logging
from datetime import datetime

import nltk
import streamlit as st

nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')

nlp = spacy.load("en_core_web_sm")

# Configuración del registro
logging.basicConfig(filename='logs/document_analysis.log', level=logging.DEBUG)

lemmatizer = WordNetLemmatizer()

# Define las funciones aquí...
def extraer_texto_docx(docx_file):
    texto = ""
    doc = Document(docx_file)
    for paragraph in doc.paragraphs:
        texto += paragraph.text + "\n"
    return texto.strip()

def extraer_texto_pdf(pdf_file):
    return pdfminer.high_level.extract_text(pdf_file)

def leer_archivo_texto(txt_file):
    return txt_file.read().decode('utf-8')

def procesar_texto(texto):
    return nlp(texto)

def tokenizar_lineamientos(lineamientos):
    tokens = []
    for lineamiento in lineamientos:
        doc = nlp(lineamiento)
        tokens.extend([token.text for token in doc])
    return list(set(tokens))

def vectorizar_texto(texto, tokens_referencia):
    vectorizer = TfidfVectorizer(vocabulary=tokens_referencia, lowercase=False)
    vector_tfidf = vectorizer.fit_transform([texto])
    return vector_tfidf.toarray()

def guardar_diccionario_en_csv(diccionario, ruta_archivo_csv):
    with open(ruta_archivo_csv, 'w', encoding='utf-8', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(["Regla", "Vector"])
        for regla, vector in diccionario.items():
            writer.writerow([regla, json.dumps(vector)])

def cargar_diccionario_desde_csv(ruta_archivo):
    diccionario = {}
    with open(ruta_archivo, 'r', encoding='utf-8') as archivo_csv:
        reader = csv.reader(archivo_csv)
        next(reader)
        for row in reader:
            if len(row) >= 2:
                key = row[0]
                value = json.loads(row[1])
                diccionario[key] = value
    return diccionario

def encontrar_diferencias(documento1, documento2):
    diferencias = []
    try:
        if isinstance(documento1, str) and isinstance(documento2, str):
            lineas1 = documento1.split('\n')
            lineas2 = documento2.split('\n')
            for i, (linea1, linea2) in enumerate(zip(lineas1, lineas2), start=1):
                if linea1 != linea2:
                    diferencias.append((linea1, linea2, i, "Línea"))
        else:
            for i, parrafo1 in enumerate(documento1.paragraphs, start=1):
                if i <= len(documento2.paragraphs):
                    parrafo2 = documento2.paragraphs[i-1]
                    if parrafo1.text != parrafo2.text:
                        diferencias.append((parrafo1.text, parrafo2.text, i, "Párrafo"))
                else:
                    diferencias.append((parrafo1.text, "", i, "Párrafo"))

        return diferencias
    except Exception as e:
        logging.error(f"Error al encontrar diferencias: {e}")

def vectorizar_y_tokenizar_diferencias(diferencias, tokens_referencia, nombre_documento_comparar, nombre_documento_referencia):
    diferencias_vectorizadas = []
    for diferencia in diferencias:
        texto_diferencia = diferencia[0] + " " + diferencia[1]
        tokens_diferencia = tokenizar_lineamientos([texto_diferencia])
        vector_tfidf_diferencia = vectorizar_texto(texto_diferencia, tokens_referencia)
        diferencias_vectorizadas.append({
            "Texto Diferencia": texto_diferencia,
            "Vector": vector_tfidf_diferencia.tolist()[0]
        })
    if not diferencias_vectorizadas:
        return None
    df_diferencias = pd.DataFrame(diferencias_vectorizadas)
    ruta_directorio = "data/output/"
    os.makedirs(ruta_directorio, exist_ok=True)
    nombre_archivo_csv = f"{nombre_documento_referencia}_{nombre_documento_comparar}_diferencias.csv"
    ruta_archivo_csv = os.path.join(ruta_directorio, nombre_archivo_csv)
    df_diferencias.to_csv(ruta_archivo_csv, index=False, encoding='utf-8')
    return diferencias_vectorizadas

def almacenar_reglas_vectorizadas(texto_manual, tokens_referencia):
    reglas_vectorizadas = {}
    reglas = texto_manual.split("\n")
    for regla in reglas:
        regla = regla.strip()
        if regla:
            vector_tfidf = vectorizar_texto(regla, tokens_referencia)
            reglas_vectorizadas[regla] = vector_tfidf.tolist()[0]
    ruta_archivo_csv = "data/output/reglas_vectorizadas.csv"
    guardar_diccionario_en_csv(reglas_vectorizadas, ruta_archivo_csv)
    return reglas_vectorizadas

def cargar_y_vectorizar_manual(file, file_type, tokens_referencia):
    if file_type == "pdf":
        texto_manual = extraer_texto_pdf(file)
    elif file_type == "txt":
        texto_manual = leer_archivo_texto(file)
    elif file_type == "docx":
        texto_manual = extraer_texto_docx(file)
    else:
        return None

    reglas_vectorizadas = almacenar_reglas_vectorizadas(texto_manual, tokens_referencia)
    ruta_archivo_csv = "data/output/manual_vectorizado.csv"
    guardar_diccionario_en_csv(reglas_vectorizadas, ruta_archivo_csv)
    return ruta_archivo_csv
